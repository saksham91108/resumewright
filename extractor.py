import os
import pdfplumber
import pytesseract
from pdf2image import convert_from_path
from docx import Document

MIN_TEXT_LENGTH = 50  # tweak based on testing


class ExtractionError(Exception):
    pass


def extract_text_direct(pdf_path):
    """Try to extract embedded text directly from a PDF."""
    text_parts = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_ocr(pdf_path):
    """Fallback: rasterize PDF pages to images and run OCR."""
    images = convert_from_path(pdf_path)
    text_parts = []
    for img in images:
        text_parts.append(pytesseract.image_to_string(img))
    return "\n".join(text_parts)


def extract_text_from_pdf(path):
    try:
        text = extract_text_direct(path)
    except Exception as e:
        raise ExtractionError(f"Failed to open/parse PDF: {e}")

    if len(text.strip()) < MIN_TEXT_LENGTH:
        print("[extractor] Direct extraction too short — falling back to OCR")
        try:
            text = extract_text_ocr(path)
        except Exception as e:
            raise ExtractionError(f"OCR fallback also failed: {e}")

    return text


def extract_text_from_docx(path):
    """Extract text from a .docx file, including paragraphs and table cells."""
    try:
        doc = Document(path)
    except Exception as e:
        raise ExtractionError(f"Failed to open/parse DOCX: {e}")

    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def extract_text_from_txt(path):
    """Read a plain-text resume file, trying a couple of common encodings."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            with open(path, "r", encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
        except Exception as e:
            raise ExtractionError(f"Failed to read TXT file: {e}")
    raise ExtractionError("Failed to read TXT file: could not decode with common encodings.")


def extract_text(path):
    """
    Main entry point: dispatches to the right extractor based on file extension.
    Supports .pdf, .docx, and .txt.
    """
    if not os.path.exists(path):
        raise ExtractionError(f"File not found: {path}")

    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext == ".docx":
        return extract_text_from_docx(path)
    elif ext == ".txt":
        return extract_text_from_txt(path)
    else:
        raise ExtractionError(
            f"Unsupported file type: '{ext}'. Please upload a .pdf, .docx, or .txt resume."
        )


if __name__ == "__main__":
    import sys
    path = sys.argv[1]
    result = extract_text(path)
    print("--- EXTRACTED TEXT ---")
    print(result)
    print(f"\n--- LENGTH: {len(result)} chars ---")